import os
import re
from PyPDF2 import PdfReader
from docx import Document

class ResumeParser:
    """Parse resume files and extract key information"""
    
    def __init__(self, file_path):
        self.file_path = file_path
        self.content = ""
        self.file_type = os.path.splitext(file_path)[1].lower()
        
    def parse(self):
        """Parse the resume based on file type"""
        if self.file_type == '.pdf':
            self._parse_pdf()
        elif self.file_type in ['.docx', '.doc']:
            self._parse_docx()
        elif self.file_type == '.txt':
            self._parse_txt()
        else:
            raise ValueError(f"Unsupported file type: {self.file_type}")
        
        return self.extract_information()
    
    def _parse_pdf(self):
        """Extract text from PDF"""
        try:
            with open(self.file_path, 'rb') as file:
                reader = PdfReader(file)
                for page in reader.pages:
                    self.content += page.extract_text()
        except Exception as e:
            raise Exception(f"Error parsing PDF: {str(e)}")
    
    def _parse_docx(self):
        """Extract text from DOCX"""
        try:
            doc = Document(self.file_path)
            for para in doc.paragraphs:
                self.content += para.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        self.content += cell.text + "\n"
        except Exception as e:
            raise Exception(f"Error parsing DOCX: {str(e)}")
    
    def _parse_txt(self):
        """Extract text from TXT"""
        try:
            with open(self.file_path, 'r', encoding='utf-8') as file:
                self.content = file.read()
        except Exception as e:
            raise Exception(f"Error parsing TXT: {str(e)}")
    
    def extract_information(self):
        """Extract structured information from resume"""
        resume_data = {
            'full_text': self.content,
            'skills': self._extract_skills(),
            'experience': self._extract_experience(),
            'education': self._extract_education(),
            'contact_info': self._extract_contact_info(),
            'certifications': self._extract_certifications(),
            'key_highlights': self._extract_highlights()
        }
        return resume_data
    
    def _extract_skills(self):
        """Extract technical and professional skills"""
        skills = []
        skill_keywords = [
            'python', 'java', 'javascript', 'typescript', 'react', 'angular', 'vue',
            'nodejs', 'express', 'django', 'flask', 'sql', 'mongodb', 'postgres',
            'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'git', 'ci/cd',
            'machine learning', 'deep learning', 'nlp', 'computer vision',
            'api', 'rest', 'graphql', 'agile', 'scrum', 'project management',
            'leadership', 'communication', 'team collaboration', 'problem solving',
            'c++', 'c#', '.net', 'golang', 'rust', 'scala', 'ruby', 'php'
        ]
        
        content_lower = self.content.lower()
        for keyword in skill_keywords:
            if keyword in content_lower:
                skills.append(keyword)
        
        return list(set(skills))
    
    def _extract_experience(self):
        """Extract work experience"""
        experience = []
        # Look for common patterns: job titles and companies
        lines = self.content.split('\n')
        for i, line in enumerate(lines):
            if any(keyword in line.lower() for keyword in ['engineer', 'developer', 'manager', 'analyst', 'designer', 'architect']):
                experience.append(line.strip())
        
        return experience[:5]  # Return top 5 experiences
    
    def _extract_education(self):
        """Extract education information"""
        education = []
        degree_keywords = ['bachelor', 'master', 'phd', 'diploma', 'bs', 'ms', 'btech', 'mtech', 'b.tech', 'm.tech']
        
        lines = self.content.split('\n')
        for line in lines:
            if any(keyword in line.lower() for keyword in degree_keywords):
                education.append(line.strip())
        
        return education
    
    def _extract_contact_info(self):
        """Extract contact information"""
        contact = {}
        
        # Email pattern
        email_pattern = r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}'
        emails = re.findall(email_pattern, self.content)
        if emails:
            contact['email'] = emails[0]
        
        # Phone pattern
        phone_pattern = r'(\+\d{1,3}[-.\s]?)?\d{3}[-.\s]?\d{3}[-.\s]?\d{4}'
        phones = re.findall(phone_pattern, self.content)
        if phones:
            contact['phone'] = phones[0]
        
        return contact
    
    def _extract_certifications(self):
        """Extract certifications"""
        certifications = []
        cert_keywords = ['certified', 'certification', 'aws', 'azure', 'google', 'pmp', 'scrum', 'cissp', 'comptia']
        
        lines = self.content.split('\n')
        for line in lines:
            if any(keyword in line.lower() for keyword in cert_keywords):
                certifications.append(line.strip())
        
        return certifications
    
    def _extract_highlights(self):
        """Extract key accomplishments and highlights"""
        highlights = []
        achievement_keywords = ['accomplished', 'achieved', 'improved', 'increased', 'led', 'developed', 'designed', 'created']
        
        lines = self.content.split('\n')
        for line in lines:
            if any(keyword in line.lower() for keyword in achievement_keywords) and len(line) > 20:
                highlights.append(line.strip())
        
        return highlights[:5]
